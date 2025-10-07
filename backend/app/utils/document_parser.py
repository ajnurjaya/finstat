import os
from pathlib import Path
from typing import Dict, List
from docx import Document
from docling.document_converter import DocumentConverter


class DocumentParser:
    """Parse different document formats and extract text content"""

    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, any]:
        """
        Extract text from PDF using Docling (IBM's advanced PDF parser)
        Handles complex layouts, tables, headers, footers, and multi-column text
        """
        try:
            print(f"📄 Parsing PDF with Docling: {Path(file_path).name}")

            # Initialize Docling converter
            converter = DocumentConverter()

            # Convert PDF to document
            result = converter.convert(file_path)

            # Extract text from document
            # Docling preserves document structure, tables, and formatting
            full_text = result.document.export_to_markdown()

            # Also get page-by-page breakdown if available
            pages_content = []
            page_count = 0

            # Docling returns structured content - extract pages
            for item in result.document.iterate_items():
                if hasattr(item, 'prov'):
                    page_num = item.prov[0].page if item.prov else page_count + 1
                    if page_num > page_count:
                        page_count = page_num

            # Get tables separately for better structure
            tables = []
            for table in result.document.tables:
                # Pass doc argument to avoid deprecation warning
                try:
                    table_data = table.export_to_dataframe(doc=result.document)
                except:
                    table_data = str(table)

                tables.append({
                    "data": table_data,
                    "caption": getattr(table, 'caption', '')
                })

            print(f"✅ Docling parsed: {page_count} pages, {len(tables)} tables")

            return {
                "success": True,
                "text": full_text,
                "page_count": page_count,
                "tables_found": len(tables),
                "tables": tables,
                "format": "pdf",
                "parser": "docling"
            }

        except Exception as e:
            print(f"❌ Docling parsing failed: {str(e)}")
            return {
                "success": False,
                "error": f"PDF parsing failed with Docling: {str(e)}",
                "format": "pdf"
            }

    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, any]:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # Also extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)

            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n=== TABLES ===\n\n" + "\n".join(tables_text)

            return {
                "success": True,
                "text": full_text,
                "paragraphs": paragraphs,
                "tables_found": len(doc.tables),
                "format": "docx"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX parsing failed: {str(e)}",
                "format": "docx"
            }

    @staticmethod
    def parse_txt(file_path: str) -> Dict[str, any]:
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()

            return {
                "success": True,
                "text": text.strip(),
                "format": "txt"
            }
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()

                return {
                    "success": True,
                    "text": text.strip(),
                    "format": "txt",
                    "encoding": "latin-1"
                }
            except Exception as e:
                return {
                    "success": False,
                    "error": f"TXT parsing failed: {str(e)}",
                    "format": "txt"
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"TXT parsing failed: {str(e)}",
                "format": "txt"
            }

    @staticmethod
    def parse_document(file_path: str) -> Dict[str, any]:
        """Main method to parse any supported document format"""
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.pdf':
            return DocumentParser.parse_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return DocumentParser.parse_docx(file_path)
        elif file_extension == '.txt':
            return DocumentParser.parse_txt(file_path)
        else:
            return {
                "success": False,
                "error": f"Unsupported file format: {file_extension}",
                "format": file_extension
            }
